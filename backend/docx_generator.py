import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree
import xml.sax.saxutils as saxutils

def set_column_count(section, count):
    if count == 1:
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')
        if cols:
            sectPr.remove(cols[0])
        return

    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if cols:
        cols_element = cols[0]
    else:
        cols_element = OxmlElement('w:cols')
        sectPr.append(cols_element)
    cols_element.set(qn('w:num'), str(count))
    cols_element.set(qn('w:space'), '708') # standard spacing between columns

def add_figure_par(doc, figure, base_font):
    filepath = figure.get("filepath")
    if not filepath or not os.path.exists(filepath):
        return
        
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=Inches(3.3))
        
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"Fig. {figure.get('figure_number', '')}. {figure.get('caption', '')}")
        cap_run.font.name = base_font
        cap_run.font.size = Pt(8)
    except Exception as e:
        print(f"Error adding figure: {e}")

def add_table_par(doc, tab_data, base_font):
    try:
        data = tab_data.get("data", [])
        if not data: return
        
        # Caption above table
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(f"TABLE {tab_data.get('number', '')}\n")
        cap_run.font.name = base_font
        cap_run.font.size = Pt(8)
        cap_run.small_caps = True
        
        cap_run_b = cap_p.add_run(tab_data.get("caption", "").upper())
        cap_run_b.font.name = base_font
        cap_run_b.font.size = Pt(8)
        
        rows = len(data)
        cols = len(data[0]) if rows > 0 else 0
        t = doc.add_table(rows=rows, cols=cols)
        t.style = 'Table Grid'
        
        for r_idx, row in enumerate(data):
            for c_idx, val in enumerate(row):
                cell = t.cell(r_idx, c_idx)
                cell.text = str(val)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = base_font
                        run.font.size = Pt(9)
    except Exception as e:
        print(f"Error adding table: {e}")

def get_figure_by_num(figures, num):
    for f in figures:
        if str(f.get("figure_number", "")).strip() == str(num).strip():
            return f
    return None

def get_table_by_num(tables, num):
    for t in tables:
        if str(t.get("number", "")).strip() == str(num).strip():
            return t
    return None

def process_text_runs(p, text, font_name):
    # Process text for math equations wrapped in [MATH]...[/MATH]
    # And normal text otherwise
    parts = re.split(r'(\[MATH\].*?\[/MATH\])', text, flags=re.DOTALL)
    
    for part in parts:
        if part.startswith("[MATH]") and part.endswith("[/MATH]"):
            math_text = part[6:-7].strip()
            math_text_escaped = saxutils.escape(math_text)
            
            run = p.add_run()
            # Try appending a native Word Equation Element (OMML)
            math_xml = f'''
            <m:oMath xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
                     xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">
               <m:r>
                  <w:rPr>
                     <w:rFonts w:ascii="Cambria Math" w:hAnsi="Cambria Math"/>
                     <w:i/>
                  </w:rPr>
                  <m:t>{math_text_escaped}</m:t>
               </m:r>
            </m:oMath>
            '''
            try:
                math_element = etree.fromstring(math_xml)
                p._p.append(math_element)
            except Exception as e:
                # Fallback
                run = p.add_run(math_text)
                run.font.name = "Cambria Math"
                run.italic = True
        else:
            if part:
                p.add_run(part)

def generate_docx(config: dict, sections: list, figures: list, tables: list, output_path: str):
    doc = Document()
    
    # 1. Setup Base styles
    font_name = config.get("font", "Times New Roman")
    style = doc.styles['Normal']
    font = style.font
    font.name = font_name
    font.size = Pt(10)
    
    margins = config.get("margins", {})
    t_m = margins.get("top", 0.75)
    b_m = margins.get("bottom", 1.0)
    l_m = margins.get("left", 0.63)
    r_m = margins.get("right", 0.63)
    
    sec1 = doc.sections[0]
    sec1.top_margin = Inches(t_m)
    sec1.bottom_margin = Inches(b_m)
    sec1.left_margin = Inches(l_m)
    sec1.right_margin = Inches(r_m)
    
    # Determine base layout width
    is_ieee = config.get("preset") == "ieee"
    set_column_count(sec1, 1) # Start single column for headers
    
    body_section_started = False
    
    # Render loop
    for sec_data in sections:
        sec_name = sec_data.name.strip()
        content = sec_data.content.strip()
        
        # In custom preset or IEEE, handle Title explicitly
        if sec_name.lower() == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.bold = True
            run.font.size = Pt(config.get("title_size", 24))
            continue
            
        if sec_name.lower() == "authors":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.font.size = Pt(11)
            continue
            
        if sec_name.lower() == "affiliations":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(content)
            run.italic = True
            continue
            
        # Start double column body layout for IEEE preset if not yet started
        if is_ieee and not body_section_started:
            new_sec = doc.add_section(WD_SECTION.CONTINUOUS)
            new_sec.top_margin = Inches(t_m)
            new_sec.bottom_margin = Inches(b_m)
            new_sec.left_margin = Inches(l_m)
            new_sec.right_margin = Inches(r_m)
            set_column_count(new_sec, 2)
            body_section_started = True
            
        # Section Header Rendering
        if content:
            p = doc.add_paragraph()
            if sec_name.lower() == "abstract" and is_ieee:
                run = p.add_run("Abstract—")
                run.bold = True
                run.italic = True
            elif sec_name.lower() == "keywords" and is_ieee:
                run = p.add_run("Keywords—")
                run.bold = True
                run.italic = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(sec_name.upper() if is_ieee else sec_name)
                run.font.name = font_name
                run.font.size = Pt(10)
                p = doc.add_paragraph() # Prep next par for content
            
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Content Rendering
            # Looking for figures embedded in text via [FIGURE X]
            lines = content.split('\n')
            for line in lines:
                if not line.strip():
                    continue
                    
                while True:
                    match_fig = re.search(r'\[FIGURE\s+(\w+)\]', line)
                    match_tab = re.search(r'\[TABLE\s+(\w+)\]', line)
                    
                    match = None
                    is_fig = False
                    if match_fig and match_tab:
                        if match_fig.start() < match_tab.start():
                            match, is_fig = match_fig, True
                        else:
                            match, is_fig = match_tab, False
                    elif match_fig:
                        match, is_fig = match_fig, True
                    elif match_tab:
                        match, is_fig = match_tab, False

                    if not match:
                        process_text_runs(p, line, font_name)
                        break

                    # Render text before tag
                    before = line[:match.start()]
                    if before.strip():
                        process_text_runs(p, before, font_name)

                    if is_fig:
                        fig_num = match.group(1)
                        fig_data = get_figure_by_num(figures, fig_num)
                        if fig_data:
                            add_figure_par(doc, fig_data, font_name)
                            fig_data['_rendered'] = True
                    else:
                        tab_num = match.group(1)
                        tab_data = get_table_by_num(tables, tab_num)
                        if tab_data:
                            add_table_par(doc, tab_data, font_name)
                            tab_data['_rendered'] = True

                    # Prepare next paragraph for text after tag
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    line = line[match.end():]
                    
                # End of line paragraphing
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
    # Render any unrendered figures at bottom
    for fig in figures:
        if not fig.get('_rendered'):
            add_figure_par(doc, fig, font_name)

    # Render any unrendered tables at bottom
    for tab in tables:
        if not tab.get('_rendered'):
            add_table_par(doc, tab, font_name)
            
    doc.save(output_path)
